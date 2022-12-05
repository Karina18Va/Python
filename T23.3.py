import docx
import re

d = docx.Document('Word.docx')
for a in range(len(d.paragraphs)):
    for t in range(len(d.paragraphs[a].runs)):
        if re.match('^M{0,3}(CM|CD|D?C{0,3})?(XC|XL|L?X{0,3})?(IX|IV|V?I{0,3})?$', d.paragraphs[a].runs[t].text):
            d.paragraphs[a].runs[t].font.color.rgb = docx.shared.RGBColor(0, 255, 0)
            d.paragraphs[a].runs[t].bold = True
            print('Yeas')
        elif re.match('[MCDXLIV]+', d.paragraphs[a].runs[t].text):
            d.paragraphs[a].runs[t].font.color.rgb = docx.shared.RGBColor(255, 0, 0)
            d.paragraphs[a].runs[t].bold = True
            print('No')

d.save('Word2.docx')